"""
generate_test_docx.py
Creates a rich .docx file with complex tables (merged cells, colored headers),
nested lists, inline formatting, and other structures that force Pandoc to emit
HTML table tags — ideal for testing the extraction pipeline.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = Path("test-docs/complex_test_document.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def bold_colored(para, text, r=0, g=0, b=0, size=11, bold=True):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(r, g, b)
    return run

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading("Product Specifications & Lifecycle Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph(
    "This document is generated specifically to test the Pandoc extraction pipeline. "
    "It contains complex multi-level tables with merged cells and coloured headers, "
    "nested bullet and numbered lists, inline text formatting, and multi-paragraph cells — "
    "all structures that cause Pandoc to emit HTML table tags."
)

# ── Section 1: Simple text formatting ────────────────────────────────────────
doc.add_heading("1. Inline Text Formatting", level=1)
p = doc.add_paragraph()
bold_colored(p, "Bold red text.  ", 200, 0, 0)
run = p.add_run("Italic text.  ")
run.italic = True
run = p.add_run("Underlined text.  ")
run.underline = True
run = p.add_run("Bold + Italic + Blue.")
run.bold = True; run.italic = True
run.font.color.rgb = RGBColor(0, 70, 180)

# ── Section 2: Nested lists ──────────────────────────────────────────────────
doc.add_heading("2. Nested Bullet & Numbered Lists", level=1)
styles = doc.styles
for item in ["Alpha Product Line", "Beta Product Line", "Gamma Product Line"]:
    doc.add_paragraph(item, style="List Bullet")
    for sub in ["Sub-category A", "Sub-category B"]:
        doc.add_paragraph(sub, style="List Bullet 2")
        doc.add_paragraph("Detail level — specifications here", style="List Bullet 3")

doc.add_heading("Numbered Steps", level=2)
for i, step in enumerate([
    "Gather all input DOCX files.",
    "Run them through the Pandoc pipeline.",
    "Review the generated Markdown output.",
    "Verify with the QA team.",
], 1):
    doc.add_paragraph(step, style="List Number")

# ── Section 3: Simple 4-col table (standard MD table) ───────────────────────
doc.add_heading("3. Simple Product Table", level=1)
doc.add_paragraph("This is a standard table Pandoc should render as a GFM pipe table:")
tbl = doc.add_table(rows=5, cols=4)
tbl.style = "Table Grid"
headers = ["Product ID", "Name", "Category", "Price (£)"]
data = [
    ["P-001", "Funeral Plan Gold",   "Insurance",   "£4,500"],
    ["P-002", "Funeral Plan Silver", "Insurance",   "£3,200"],
    ["P-003", "Probate Service",     "Legal",       "£1,800"],
    ["P-004", "Will Writing Kit",    "Legal",       "£299"],
]
for j, h in enumerate(headers):
    cell = tbl.rows[0].cells[j]
    cell.text = ""
    set_cell_bg(cell, "1F497D")
    bold_colored(cell.paragraphs[0], h, 255, 255, 255, size=11)

for i, row_data in enumerate(data, 1):
    for j, val in enumerate(row_data):
        tbl.rows[i].cells[j].text = val

doc.add_paragraph()

# ── Section 4: Complex merged-cell table (forces HTML output) ────────────────
doc.add_heading("4. Complex Merged-Cell Table (triggers HTML output)", level=1)
doc.add_paragraph(
    "This table uses cell merging (colspan/rowspan). Pandoc cannot represent this "
    "in standard Markdown, so it will emit HTML <table> tags."
)

complex_tbl = doc.add_table(rows=6, cols=5)
complex_tbl.style = "Table Grid"

# Row 0: Fully merged header across all 5 cols
header_row = complex_tbl.rows[0]
for cell in header_row.cells:
    set_cell_bg(cell, "C00000")
a = header_row.cells[0].merge(header_row.cells[4])
a.text = ""
bold_colored(a.paragraphs[0], "ANNUAL PRODUCT LIFECYCLE SUMMARY", 255, 255, 255, size=13)
a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Row 1: Two merged groups + one standalone
r1 = complex_tbl.rows[1]
for cell in r1.cells:
    set_cell_bg(cell, "4472C4")

g1 = r1.cells[0].merge(r1.cells[1])
g1.text = ""
bold_colored(g1.paragraphs[0], "Product Details", 255, 255, 255)

g2 = r1.cells[2].merge(r1.cells[3])
g2.text = ""
bold_colored(g2.paragraphs[0], "Financial Performance", 255, 255, 255)

r1.cells[4].text = ""
bold_colored(r1.cells[4].paragraphs[0], "Status", 255, 255, 255)

# Row 2: Sub-headers
r2 = complex_tbl.rows[2]
for cell in r2.cells:
    set_cell_bg(cell, "D9E1F2")
for j, h in enumerate(["Code", "Name", "Q1 Revenue", "Q2 Revenue", "Lifecycle"]):
    bold_colored(r2.cells[j].paragraphs[0], h, 31, 73, 125)

# Rows 3-5: Data + one row-merged cell in column 4
row_data = [
    ["P-001", "Funeral Plan Gold",   "£420,000", "£380,000"],
    ["P-002", "Funeral Plan Silver", "£290,000", "£310,000"],
    ["P-003", "Probate Service",     "£120,000", "£145,000"],
]
status_labels = ["Active", "Active", "Under Review"]
status_colors = ["70AD47", "70AD47", "ED7D31"]

for i, (rd, sl, sc) in enumerate(zip(row_data, status_labels, status_colors), 3):
    row = complex_tbl.rows[i]
    for j, val in enumerate(rd):
        row.cells[j].text = val
    set_cell_bg(row.cells[4], sc)
    bold_colored(row.cells[4].paragraphs[0], sl, 255, 255, 255)

doc.add_paragraph()

# ── Section 5: Table with multi-paragraph cells ───────────────────────────────
doc.add_heading("5. Table With Multi-Paragraph Cells", level=1)
doc.add_paragraph(
    "Cells with multiple paragraphs also force Pandoc into HTML mode "
    "because GFM pipe tables cannot contain newlines inside a cell."
)
mp_tbl = doc.add_table(rows=3, cols=3)
mp_tbl.style = "Table Grid"

header_labels = ["Feature", "Description", "Notes"]
for j, h in enumerate(header_labels):
    set_cell_bg(mp_tbl.rows[0].cells[j], "375623")
    bold_colored(mp_tbl.rows[0].cells[j].paragraphs[0], h, 255, 255, 255)

# Row 1 — multi-para cell
mp_tbl.rows[1].cells[0].text = "Extraction Engine"
desc_cell = mp_tbl.rows[1].cells[1]
desc_cell.paragraphs[0].text = "Primary: Pandoc (gfm)"
desc_cell.add_paragraph("Fallback: MarkItDown")
desc_cell.add_paragraph("Legacy: Docling")
mp_tbl.rows[1].cells[2].text = "Configured per router"

# Row 2 — normal
mp_tbl.rows[2].cells[0].text = "Output Format"
mp_tbl.rows[2].cells[1].text = "GitHub Flavoured Markdown (.md)"
mp_tbl.rows[2].cells[2].text = "Stored in SharePoint"

doc.add_paragraph()

# ── Section 6: Blockquote-style callout ──────────────────────────────────────
doc.add_heading("6. Important Notes", level=1)
note = doc.add_paragraph(style="Intense Quote")
note.add_run(
    "All documents processed by this pipeline are subject to the 3-stage "
    "lifecycle: Raw → Draft → Final. No document reaches downstream AI systems "
    "until it has been signed off by an authorised verifier."
)

# ── Section 7: Wide table with many columns (forces HTML) ────────────────────
doc.add_heading("7. Wide Multi-Column Data Table (forces HTML)", level=1)
doc.add_paragraph(
    "A table with 8+ narrow columns where the header row spans 2 levels also "
    "cannot be expressed in GFM — Pandoc will emit HTML."
)
wide_tbl = doc.add_table(rows=5, cols=8)
wide_tbl.style = "Table Grid"

# Merged header row
wr0 = wide_tbl.rows[0]
for cell in wr0.cells:
    set_cell_bg(cell, "7030A0")
g = wr0.cells[0].merge(wr0.cells[3])
g.text = ""
bold_colored(g.paragraphs[0], "Q1 Metrics", 255, 255, 255)
g.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
g2 = wr0.cells[4].merge(wr0.cells[7])
g2.text = ""
bold_colored(g2.paragraphs[0], "Q2 Metrics", 255, 255, 255)
g2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

sub_headers = ["Jan", "Feb", "Mar", "Total Q1", "Apr", "May", "Jun", "Total Q2"]
wr1 = wide_tbl.rows[1]
for j, h in enumerate(sub_headers):
    set_cell_bg(wr1.cells[j], "C4A0DC")
    bold_colored(wr1.cells[j].paragraphs[0], h, 50, 0, 80)

data_rows = [
    ["120", "135", "142", "397",  "150", "160", "175", "485"],
    ["200", "210", "195", "605",  "220", "230", "215", "665"],
    ["80",  "90",  "85",  "255",  "100", "110", "95",  "305"],
]
for i, rd in enumerate(data_rows, 2):
    for j, val in enumerate(rd):
        wide_tbl.rows[i].cells[j].text = val

# ── Done ─────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"[OK] Document saved to: {OUT}")
print("   Upload it to the Pandoc pipeline and observe the HTML table tags.")
